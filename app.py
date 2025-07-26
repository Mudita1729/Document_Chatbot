import streamlit as st
import os
import json
from pathlib import Path
from typing import List, Dict, Any, Tuple
import hashlib
import re
from datetime import datetime

# Document processing 
from docx import Document
import PyPDF2
import pandas as pd

# Vector search and embeddings
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np

# LLM integration
from groq import Groq
import tiktoken

# --- Configuration (remains the same) ---
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
MAX_CONTEXT_LENGTH = 3000
SIMILARITY_THRESHOLD = 0.3


# --- Backend Classes (Unaltered with one minor prompt improvement) ---

class DocumentProcessor:
    """Handles document loading and text extraction"""
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return '\n'.join(text)
        except Exception as e:
            st.error(f"Error reading {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text.append(page_text)
            return '\n'.join(text)
        except Exception as e:
            st.error(f"Error reading {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            st.error(f"Error reading {file_path}: {str(e)}")
            return ""

class TextChunker:
    """Handles text chunking for better retrieval"""
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
        if not text.strip():
            return []
        words = text.split()
        chunks = [' '.join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size - overlap)]
        return [chunk.strip() for chunk in chunks if chunk.strip()]

class VectorStore:
    """Handles vector storage and similarity search"""
    
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        try:
            self.model = SentenceTransformer(model_name)
            self.index = None
            self.chunks = []
            self.chunk_metadata = []
        except Exception as e:
            st.error(f"Error loading embedding model: {str(e)}")
            raise
    
    def add_documents(self, documents: List[Dict[str, Any]]):
        try:
            all_chunks = []
            all_metadata = []
            
            for doc in documents:
                chunks = TextChunker.chunk_text(doc['content'])
                for i, chunk in enumerate(chunks):
                    all_chunks.append(chunk)
                    all_metadata.append({
                        'filename': doc['filename'],
                        'chunk_id': i,
                        'content': chunk
                    })
            
            if all_chunks:
                embeddings = self.model.encode(all_chunks, show_progress_bar=False)
                dimension = embeddings.shape[1]
                self.index = faiss.IndexFlatIP(dimension)
                faiss.normalize_L2(embeddings)
                self.index.add(embeddings.astype('float32'))
                self.chunks = all_chunks
                self.chunk_metadata = all_metadata
        except Exception as e:
            st.error(f"Error adding documents to vector store: {str(e)}")
            raise
    
    def search(self, query: str, top_k: int = 10) -> List[Dict[str, Any]]:
        if not self.index or not self.chunks:
            return []
        try:
            query_embedding = self.model.encode([query])
            faiss.normalize_L2(query_embedding)
            scores, indices = self.index.search(query_embedding.astype('float32'), top_k)
            
            results = [
                {
                    'content': self.chunk_metadata[idx]['content'],
                    'filename': self.chunk_metadata[idx]['filename'],
                    'score': float(score),
                    'chunk_id': self.chunk_metadata[idx]['chunk_id']
                }
                for score, idx in zip(scores[0], indices[0])
                if score > SIMILARITY_THRESHOLD and idx < len(self.chunk_metadata)
            ]
            return results
        except Exception as e:
            st.error(f"Error searching vector store: {str(e)}")
            return []

class GroqLLM:
    """Handles interaction with Groq API"""
    
    def __init__(self, api_key: str):
        try:
            self.client = Groq(api_key=api_key)
            self.model = "llama3-8b-8192"
        except Exception as e:
            st.error(f"Error initializing Groq client: {str(e)}")
            raise
    
    def count_tokens(self, text: str) -> int:
        return len(text.split()) * 1.3
    
    def generate_answer(self, question: str, context_chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        try:
            context_parts = []
            sources = set()
            
            for chunk in context_chunks[:5]:
                context_parts.append(f"[Source: {chunk['filename']}]\n{chunk['content']}")
                sources.add(chunk['filename'])
            
            context = "\n\n".join(context_parts)
            
            if self.count_tokens(context) > MAX_CONTEXT_LENGTH:
                words = context.split()
                truncated_words = words[:int(MAX_CONTEXT_LENGTH * 0.7)]
                context = ' '.join(truncated_words)
            
            # **Slightly improved prompt for better inline citation**
            prompt = f"""You are a helpful AI assistant that answers questions based strictly on the provided document content. 

INSTRUCTIONS:
1. Use ONLY the provided document content to answer.
2. If the answer isn't in the documents, state: "I don't have sufficient information in the provided documents to answer this question."
3. After each sentence or claim that comes from a document, cite its source like this: [Source: filename.pdf].
4. Provide a clear, well-structured answer. Do not make assumptions.

DOCUMENT CONTEXT:
{context}

QUESTION: {question}

ANSWER:"""

            completion = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=1024,
                top_p=0.95
            )
            
            answer = completion.choices[0].message.content
            
            return {
                'answer': answer,
                'sources': list(sources),
                'context_chunks': context_chunks[:5]
            }
        except Exception as e:
            return {
                'answer': f"An error occurred while generating the answer: {str(e)}",
                'sources': [],
                'context_chunks': []
            }

class DocumentQASystem:
    """Main system orchestrating all components"""
    
    def __init__(self, groq_api_key: str):
        try:
            self.vector_store = VectorStore()
            self.llm = GroqLLM(groq_api_key)
            self.documents = []
        except Exception as e:
            st.error(f"Error initializing Document Q&A System: {str(e)}")
            raise
    
    def _load_and_process(self, new_documents: List[Dict[str, Any]]) -> int:
        if not new_documents:
            return 0
        
        self.documents.extend(new_documents)
        self.vector_store.add_documents(self.documents)
        return len(new_documents)

    def load_documents_from_directory(self, directory_path: str) -> int:
        supported_extensions = {'.docx', '.pdf', '.txt'}
        new_documents = []
        directory = Path(directory_path)
        if not directory.exists() or not directory.is_dir():
            st.warning(f"Directory not found or is not valid: {directory_path}")
            return 0
        
        for file_path in directory.rglob('*'):
            if file_path.suffix.lower() in supported_extensions and file_path.is_file():
                content = ""
                try:
                    if file_path.suffix.lower() == '.docx':
                        content = DocumentProcessor.extract_text_from_docx(str(file_path))
                    elif file_path.suffix.lower() == '.pdf':
                        content = DocumentProcessor.extract_text_from_pdf(str(file_path))
                    elif file_path.suffix.lower() == '.txt':
                        content = DocumentProcessor.extract_text_from_txt(str(file_path))
                    
                    if content.strip():
                        new_documents.append({
                            'filename': file_path.name,
                            'filepath': str(file_path),
                            'content': content
                        })
                except Exception as e:
                    st.warning(f"Skipping {file_path.name}: {str(e)}")
        
        return self._load_and_process(new_documents)

    def process_uploaded_files(self, uploaded_files) -> int:
        new_documents = []
        temp_dir = "temp_uploads"
        os.makedirs(temp_dir, exist_ok=True)
        
        for uploaded_file in uploaded_files:
            content = ""
            temp_path = os.path.join(temp_dir, uploaded_file.name)
            
            try:
                with open(temp_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                
                if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    content = DocumentProcessor.extract_text_from_docx(temp_path)
                elif uploaded_file.type == "application/pdf":
                    content = DocumentProcessor.extract_text_from_pdf(temp_path)
                elif uploaded_file.type == "text/plain":
                    content = DocumentProcessor.extract_text_from_txt(temp_path)
                
                if content.strip():
                    new_documents.append({
                        'filename': uploaded_file.name,
                        'filepath': uploaded_file.name,
                        'content': content
                    })
            except Exception as e:
                st.warning(f"Error processing {uploaded_file.name}: {str(e)}")
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        return self._load_and_process(new_documents)
    
    def answer_question(self, question: str) -> Dict[str, Any]:
        if not self.documents:
            return {'answer': "No documents loaded.", 'sources': [], 'context_chunks': []}
        
        try:
            relevant_chunks = self.vector_store.search(question, top_k=10)
            if not relevant_chunks:
                return {'answer': "I couldn't find any relevant information in the documents.", 'sources': [], 'context_chunks': []}
            return self.llm.generate_answer(question, relevant_chunks)
        except Exception as e:
            return {'answer': f"Error processing question: {str(e)}", 'sources': [], 'context_chunks': []}


# --- Professional Streamlit UI with FIX ---

def main():
    st.set_page_config(
        page_title="DocuMentor AI",
        page_icon="🤖",
        layout="wide"
    )

    # Custom CSS with the fix for the answer box
    st.markdown("""
    <style>
    .stSpinner > div > div {
        border-top-color: #6366F1;
    }
    .answer-container {
        background-color: #F0F2F6;
        border-left: 5px solid #6366F1;
        padding: 1rem 1rem 1rem 2rem;
        border-radius: 5px;
        margin-top: 1rem;
        color: #262730; /* <<< THIS IS THE FIX: Sets dark text color for readability */
    }
    .st-emotion-cache-1y4p8pa {
        padding-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)

    # Initialize session state
    if 'qa_system' not in st.session_state:
        st.session_state.qa_system = None
        st.session_state.documents_loaded = False
        st.session_state.last_question = ""
        st.session_state.latest_answer = None

    # --- Sidebar for Configuration and Controls ---
    with st.sidebar:
        st.title("👨‍💻 DocuMentor AI")
        st.markdown("Your personal document expert. Upload your docs, ask questions, and get instant, source-backed answers.")
        
        st.divider()

        st.header("⚙️ Step 1: API Setup")
        groq_api_key = os.getenv('GROQ_API_KEY')
        
        if groq_api_key:
            st.success("✅ Groq API Key found!", icon="✔️")
            final_api_key = groq_api_key
        else:
            final_api_key = st.text_input("Enter your Groq API Key", type="password", help="Get your key from [Groq Console](https://console.groq.com/keys)")

        if not final_api_key:
            st.warning("Please provide your Groq API key to proceed.", icon="⚠️")
            st.stop()
            
        st.divider()

        st.header("📁 Step 2: Load Documents")
        uploaded_files = st.file_uploader(
            "Upload Files",
            type=['docx', 'pdf', 'txt'],
            accept_multiple_files=True,
            help="Upload one or more documents."
        )
        directory_path = st.text_input("Or provide a Directory Path", help="e.g., C:/Users/YourUser/Documents")

        load_button = st.button("Load & Process Documents", type="primary", use_container_width=True, disabled=st.session_state.documents_loaded)
    
    # --- Main Content Area ---
    st.header("🤖 Ask Your Document-Based Questions", divider='rainbow')

    # Handling document loading
    if load_button:
        with st.spinner("Initializing AI system and processing documents... This may take a moment."):
            try:
                qa_system = DocumentQASystem(final_api_key)
                
                total_loaded = 0
                if uploaded_files:
                    total_loaded += qa_system.process_uploaded_files(uploaded_files)
                if directory_path:
                    total_loaded += qa_system.load_documents_from_directory(directory_path)
                
                if total_loaded > 0:
                    st.session_state.qa_system = qa_system
                    st.session_state.documents_loaded = True
                    st.success(f"✅ Success! {total_loaded} new document(s) loaded and ready.")
                else:
                    st.error("❌ No documents were loaded. Please check your uploads or directory path.")
            
            except Exception as e:
                st.error(f"Failed to initialize the system: {e}")

    # Displaying the main interface
    if not st.session_state.documents_loaded:
        st.info("👋 Welcome! Please load your documents using the sidebar on the left to get started.")
        
        st.subheader("How It Works")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("#### 1. Configure ⚙️")
            st.markdown("In case you want to use your own API key please enter your Groq API key in the sidebar.Otherwise default API key would be used")
        with col2:
            st.markdown("#### 2. Load Docs 📁")
            st.markdown("Upload your `.pdf`, `.docx`, or `.txt` files, or point to a local directory.")
        with col3:
            st.markdown("#### 3. Ask Away 💬")
            st.markdown("Once loaded, ask any question about your documents and get cited answers.")

    else:
        # Interface for asking questions
        doc_count = len(st.session_state.qa_system.documents)
        with st.expander(f"📄 View {doc_count} Loaded Document(s)"):
            for doc in st.session_state.qa_system.documents:
                st.write(f"- `{doc['filename']}`")

        question = st.text_input(
            "Your Question:",
            placeholder="e.g., What are the key features of product X?",
            key="question_input"
        )

        if question and question != st.session_state.last_question:
            st.session_state.last_question = question
            with st.spinner("🧠 Thinking... Searching documents and crafting your answer..."):
                try:
                    result = st.session_state.qa_system.answer_question(question)
                    st.session_state.latest_answer = result
                except Exception as e:
                    st.error(f"An error occurred: {e}")
                    st.session_state.latest_answer = None
        
        # Display the latest answer
        if st.session_state.latest_answer:
            result = st.session_state.latest_answer
            
            st.subheader("🎯 Answer")
            st.markdown(f"<div class='answer-container'>{result['answer']}</div>", unsafe_allow_html=True)

            col1, col2 = st.columns([1, 2])
            with col1:
                if result.get('sources'):
                    with st.container(border=True):
                        st.subheader("📚 Sources")
                        for source in result['sources']:
                            st.write(f"• {source}")
            
            with col2:
                 if result.get('context_chunks'):
                    with st.expander("📖 View Relevant Excerpts..."):
                        for i, chunk in enumerate(result['context_chunks']):
                            with st.container(border=True):
                                st.markdown(f"**Source:** `{chunk['filename']}` | **Relevance:** {chunk['score']:.2f}")
                                st.caption(f"Excerpt: \"{chunk['content'][:250]}...\"")

if __name__ == "__main__":
    main()